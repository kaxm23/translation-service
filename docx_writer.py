from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from typing import Dict, List, Optional, Union
import logging
from datetime import datetime
from pathlib import Path
import json

class TranslatedDocWriter:
    """
    A class for creating DOCX files from translated text with formatting.
    Created by: kaxm23
    Created on: 2025-02-09 08:28:28 UTC
    """
    
    def __init__(self,
                 template_path: Optional[str] = None,
                 font_name: str = 'Arial',
                 font_size: int = 11,
                 rtl_mode: bool = True,
                 log_level: int = logging.INFO):
        """
        Initialize the DOCX writer.
        
        Args:
            template_path: Path to template DOCX file
            font_name: Default font name
            font_size: Default font size
            rtl_mode: Enable RTL text direction for Arabic
            log_level: Logging level
        """
        # Configure logging
        logging.basicConfig(
            level=log_level,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Initialize document
        self.doc = Document(template_path) if template_path else Document()
        self.font_name = font_name
        self.font_size = font_size
        self.rtl_mode = rtl_mode
        
        # Set up default styles
        self._setup_styles()
        
        # Statistics tracking
        self.stats = {
            'paragraphs_added': 0,
            'tables_added': 0,
            'images_added': 0,
            'processing_time': None
        }

    def _setup_styles(self):
        """Configure document styles."""
        # Default paragraph style
        style = self.doc.styles.add_style('Normal_AR', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        
        # Heading styles
        for level in range(1, 4):
            style_name = f'Heading{level}_AR'
            style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = self.font_name
            font.size = Pt(self.font_size + (4 - level) * 2)
            font.bold = True
            
        # Table style
        style = self.doc.styles.add_style('Table_AR', WD_STYLE_TYPE.TABLE)
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)

    def create_translated_document(self,
                                 translation_data: Dict,
                                 include_original: bool = True,
                                 add_metadata: bool = True) -> None:
        """
        Create a DOCX document from translated text.
        
        Args:
            translation_data: Dictionary containing translation results
            include_original: Include original text
            add_metadata: Add document metadata
        """
        start_time = datetime.now()
        
        try:
            # Add metadata section
            if add_metadata and 'metadata' in translation_data:
                self._add_metadata_section(translation_data['metadata'])
            
            # Process content
            for item in translation_data.get('content', []):
                if 'type' in item:
                    if item['type'] == 'paragraph':
                        self._add_paragraph(item, include_original)
                    elif item['type'] == 'table':
                        self._add_table(item, include_original)
                    elif item['type'] == 'image':
                        self._add_image(item)
                        
            # Update statistics
            self.stats['processing_time'] = \
                (datetime.now() - start_time).total_seconds()
                
        except Exception as e:
            self.logger.error(f"Failed to create document: {str(e)}")
            raise

    def _add_metadata_section(self, metadata: Dict):
        """Add metadata section to document."""
        section = self.doc.add_section()
        
        # Add title
        title = self.doc.add_heading('Document Information', level=1)
        title.style = self.doc.styles['Heading1_AR']
        
        # Add metadata table
        table = self.doc.add_table(rows=1, cols=2)
        table.style = self.doc.styles['Table_AR']
        
        # Add metadata rows
        for key, value in metadata.items():
            row = table.add_row()
            row.cells[0].text = key.replace('_', ' ').title()
            row.cells[1].text = str(value)
            
        self.doc.add_paragraph()  # Add spacing
        self.stats['tables_added'] += 1

    def _add_paragraph(self,
                      para_data: Dict,
                      include_original: bool):
        """Add a paragraph with translation."""
        if include_original and 'original' in para_data:
            # Add original text
            orig_para = self.doc.add_paragraph()
            orig_para.style = self.doc.styles['Normal_AR']
            run = orig_para.add_run('Original: ')
            run.bold = True
            orig_para.add_run(para_data['original'])
        
        # Add translated text
        trans_para = self.doc.add_paragraph()
        trans_para.style = self.doc.styles['Normal_AR']
        
        if self.rtl_mode:
            # Set RTL paragraph direction
            pPr = trans_para._p.get_or_add_pPr()
            pPr.set(qn('w:bidi'), '1')
        
        if include_original:
            run = trans_para.add_run('Translation: ')
            run.bold = True
        
        trans_para.add_run(para_data['translated'])
        
        # Add spacing
        self.doc.add_paragraph()
        self.stats['paragraphs_added'] += 1

    def _add_table(self,
                   table_data: Dict,
                   include_original: bool):
        """Add a table with translation."""
        if include_original and 'original' in table_data:
            # Add original table
            self.doc.add_paragraph().add_run('Original Table:').bold = True
            table = self._create_table(table_data['original'])
            self.doc.add_paragraph()
        
        # Add translated table
        if include_original:
            self.doc.add_paragraph().add_run('Translated Table:').bold = True
        table = self._create_table(table_data['translated'])
        self.doc.add_paragraph()
        
        self.stats['tables_added'] += 1

    def _create_table(self, data: List[List]):
        """Create a table from data."""
        if not data or not data[0]:
            return None
            
        table = self.doc.add_table(rows=1, cols=len(data[0]))
        table.style = self.doc.styles['Table_AR']
        
        # Add header row
        for i, cell in enumerate(data[0]):
            table.rows[0].cells[i].text = str(cell)
            
        # Add data rows
        for row_data in data[1:]:
            row = table.add_row()
            for i, cell in enumerate(row_data):
                row.cells[i].text = str(cell)
        
        return table

    def _add_image(self, image_data: Dict):
        """Add an image with caption."""
        if 'path' in image_data:
            try:
                # Add image
                self.doc.add_picture(
                    image_data['path'],
                    width=Inches(6)  # Adjust width as needed
                )
                
                # Add caption
                if 'caption' in image_data:
                    caption = self.doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.add_run(image_data['caption']).italic = True
                
                self.stats['images_added'] += 1
                
            except Exception as e:
                self.logger.error(f"Failed to add image: {str(e)}")

    def save_document(self,
                     output_path: str,
                     include_stats: bool = True) -> None:
        """
        Save the document to file.
        
        Args:
            output_path: Path to save the DOCX file
            include_stats: Include processing statistics
        """
        try:
            if include_stats:
                # Add statistics section
                self.doc.add_section()
                stats_heading = self.doc.add_heading('Processing Statistics', level=1)
                stats_heading.style = self.doc.styles['Heading1_AR']
                
                stats_table = self.doc.add_table(rows=1, cols=2)
                stats_table.style = self.doc.styles['Table_AR']
                
                for key, value in self.stats.items():
                    row = stats_table.add_row()
                    row.cells[0].text = key.replace('_', ' ').title()
                    row.cells[1].text = str(value)
            
            # Save document
            self.doc.save(output_path)
            self.logger.info(f"Document saved successfully: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to save document: {str(e)}")
            raise